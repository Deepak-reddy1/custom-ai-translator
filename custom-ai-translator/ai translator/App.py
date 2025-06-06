import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from deep_translator import GoogleTranslator
from docx import Document
from fpdf import FPDF
import json


class CustomAITranslator:
    def __init__(self, root):
        self.root = root
        self.root.title("Custom AI Translator")
        self.root.geometry("800x700")
        self.root.resizable(False, False)

        # User profile
        self.user_profile = {
            "preferred_src_lang": "en",
            "preferred_tgt_lang": "es"
        }

        # Load language data
        self.languages = self.load_languages()

        # Create Gradient Header
        self.create_gradient_header()

        # Create Main Content Layout
        self.create_cards()

        # Create Footer
        self.create_footer()

        # Initialize text storage
        self.original_text = ""

    def load_languages(self):
        """Load language codes from JSON."""
        try:
            with open("languages.json", "r", encoding="utf-8") as file:
                data = json.load(file)
                return data["languages"]
        except FileNotFoundError:
            messagebox.showerror("Error", "languages.json file not found.")
            return {}

    def create_gradient_header(self):
        """Create a gradient header."""
        header_canvas = tk.Canvas(self.root, height=80, width=800)
        header_canvas.pack(fill="x")

        # Draw gradient
        gradient_color1 = "#6a11cb"  # Purple
        gradient_color2 = "#2575fc"  # Blue
        self.draw_gradient(header_canvas, gradient_color1, gradient_color2, width=800, height=80)

        # Add Title
        title_label = tk.Label(self.root, text="Custom AI Translator", font=("Helvetica", 24, "bold"), bg=gradient_color2, fg="white")
        title_label.place(x=20, y=20)

    def draw_gradient(self, canvas, color1, color2, width, height):
        """Draw a vertical gradient on a Canvas."""
        r1, g1, b1 = self.hex_to_rgb(color1)
        r2, g2, b2 = self.hex_to_rgb(color2)

        for i in range(height):
            r = int(r1 + (r2 - r1) * i / height)
            g = int(g1 + (g2 - g1) * i / height)
            b = int(b1 + (b2 - b1) * i / height)
            color = f'#{r:02x}{g:02x}{b:02x}'
            canvas.create_line(0, i, width, i, fill=color)

    def hex_to_rgb(self, hex_color):
        """Convert a HEX color to RGB."""
        hex_color = hex_color.lstrip("#")
        return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))

    def create_cards(self):
        """Create cards for Input, Translation, and Output."""
        content_frame = tk.Frame(self.root, bg="#ecf0f1")
        content_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Card 1: Input Text
        input_card = self.create_card(content_frame, "Enter Text to Translate")
        self.input_text = tk.Text(input_card, height=6, wrap=tk.WORD, relief=tk.GROOVE, bd=2)
        self.input_text.pack(fill="both", expand=True, padx=10, pady=10)

        # Card 2: Language Selection
        lang_card = self.create_card(content_frame, "Select Languages")
        lang_frame = tk.Frame(lang_card, bg="white")
        lang_frame.pack(fill="x", padx=10, pady=10)

        self.src_lang = tk.StringVar(value="English")
        src_dropdown = ttk.Combobox(lang_frame, textvariable=self.src_lang, values=list(self.languages.keys()), state="readonly")
        src_dropdown.pack(side="left", fill="x", expand=True, padx=5)
        src_dropdown.bind("<<ComboboxSelected>>", self.save_src_lang)

        self.tgt_lang = tk.StringVar(value="Spanish")
        tgt_dropdown = ttk.Combobox(lang_frame, textvariable=self.tgt_lang, values=list(self.languages.keys()), state="readonly")
        tgt_dropdown.pack(side="left", fill="x", expand=True, padx=5)
        tgt_dropdown.bind("<<ComboboxSelected>>", self.save_tgt_lang)

        # Card 3: Buttons
        button_card = self.create_card(content_frame, "Actions")
        button_frame = tk.Frame(button_card, bg="white")
        button_frame.pack(pady=10)

        translate_button = tk.Button(button_frame, text="Translate", font=("Helvetica", 12), bg="#2ecc71", fg="white", command=self.translate_text)
        translate_button.pack(side="left", padx=10)

        export_button = tk.Button(button_frame, text="Export", font=("Helvetica", 12), bg="#3498db", fg="white", command=self.export_translation)
        export_button.pack(side="left", padx=10)

        clear_button = tk.Button(button_frame, text="Clear", font=("Helvetica", 12), bg="#e74c3c", fg="white", command=self.clear_fields)
        clear_button.pack(side="left", padx=10)

        # Card 4: Translated Text Output
        output_card = self.create_card(content_frame, "Translated Text")
        self.output_text = tk.Text(output_card, height=6, wrap=tk.WORD, relief=tk.GROOVE, bd=2)
        self.output_text.pack(fill="both", expand=True, padx=10, pady=10)

    def create_footer(self):
        """Create a footer section."""
        footer_frame = tk.Frame(self.root, bg="#34495e", height=40)
        footer_frame.pack(fill="x", side="bottom")

        footer_label = tk.Label(footer_frame, text="© 2024 Custom AI Translator. All Rights Reserved.",
                                font=("Helvetica", 10), bg="#34495e", fg="white")
        footer_label.pack(pady=5)

    def create_card(self, parent, title):
        """Create a styled card frame."""
        card = tk.Frame(parent, bg="white", relief=tk.RAISED, bd=2)
        card.pack(fill="x", pady=10)

        label = tk.Label(card, text=title, font=("Helvetica", 14, "bold"), bg="white", fg="#2c3e50")
        label.pack(anchor="w", padx=10, pady=5)

        return card

    def save_src_lang(self, event):
        """Save the selected source language."""
        selected_lang = self.src_lang.get()
        if selected_lang in self.languages:
            self.user_profile["preferred_src_lang"] = self.languages[selected_lang]

    def save_tgt_lang(self, event):
        """Save the selected target language."""
        selected_lang = self.tgt_lang.get()
        if selected_lang in self.languages:
            self.user_profile["preferred_tgt_lang"] = self.languages[selected_lang]

    def translate_text(self):
        """Translate the input text."""
        src_lang_code = self.user_profile["preferred_src_lang"]
        tgt_lang_code = self.user_profile["preferred_tgt_lang"]

        text_to_translate = self.input_text.get("1.0", tk.END).strip()

        if text_to_translate:
            try:
                translated = GoogleTranslator(source=src_lang_code, target=tgt_lang_code).translate(text_to_translate)
                self.output_text.delete("1.0", tk.END)
                self.output_text.insert(tk.END, translated)

                self.original_text = text_to_translate
            except Exception as e:
                messagebox.showerror("Translation Error", str(e))
        else:
            messagebox.showwarning("Input Error", "Please enter text to translate.")

    def clear_fields(self):
        """Clear all text fields."""
        self.input_text.delete("1.0", tk.END)
        self.output_text.delete("1.0", tk.END)
        self.original_text = ""

    def export_translation(self):
        """Export the translated text."""
        translated_text = self.output_text.get("1.0", tk.END).strip()
        if not translated_text or not self.original_text:
            messagebox.showwarning("Export Error", "No translated text to export.")
            return

        file_name = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word Documents", ".docx"), ("PDF Files", ".pdf")])
        if file_name.endswith(".docx"):
            self.export_to_word(file_name, translated_text)
        elif file_name.endswith(".pdf"):
            self.export_to_pdf(file_name, translated_text)

    def export_to_word(self, file_name, translated_text):
        """Export the translation to a Word file."""
        doc = Document()
        doc.add_heading("Translations", level=1)
        doc.add_paragraph(f"Original Text ({self.user_profile['preferred_src_lang']}): {self.original_text}")
        doc.add_paragraph(f"Translated Text ({self.user_profile['preferred_tgt_lang']}): {translated_text}")
        doc.save(file_name)
        messagebox.showinfo("Export Successful", "Translated text exported to Word.")

    def export_to_pdf(self, file_name, translated_text):
        """Export the translation to a PDF file."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Translations", ln=True, align="C")
        pdf.multi_cell(0, 10, f"Original Text ({self.user_profile['preferred_src_lang']}): {self.original_text}")
        pdf.multi_cell(0, 10, f"Translated Text ({self.user_profile['preferred_tgt_lang']}): {translated_text}")
        pdf.output(file_name)
        messagebox.showinfo("Export Successful", "Translated text exported to PDF.")


if __name__ == "__main__":
    root = tk.Tk()
    app = CustomAITranslator(root)
    root.mainloop()